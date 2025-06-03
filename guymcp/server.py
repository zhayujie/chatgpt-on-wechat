from mcp.server.fastmcp import FastMCP
from dotenv import load_dotenv
import httpx
import json
import os
from bs4 import BeautifulSoup
from typing import Any
import httpx
from mcp.server.fastmcp import FastMCP
from starlette.applications import Starlette
from mcp.server.sse import SseServerTransport
from starlette.requests import Request
from starlette.routing import Mount, Route
from mcp.server import Server
import uvicorn

import os
import io
import base64
import shutil
from typing import Dict, List, Optional, Any, Union, Tuple
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from mcp.server.fastmcp import FastMCP
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import sys
from openai import OpenAI
import re  # 添加正则表达式模块

load_dotenv()

mcp = FastMCP("docs")

USER_AGENT = "docs-app/1.0"
SERPER_URL = "https://google.serper.dev/search"

docs_urls = {
    "langchain": "python.langchain.com/docs",
    "llama-index": "docs.llamaindex.ai/en/stable",
    "autogen": "microsoft.github.io/autogen/stable",
    "agno": "docs.agno.com",
    "openai-agents-sdk": "openai.github.io/openai-agents-python",
    "mcp-doc": "modelcontextprotocol.io",
    "camel-ai": "docs.camel-ai.org",
    "crew-ai": "docs.crewai.com"
}

async def search_web(query: str) -> dict | None:
    payload = json.dumps({"q": query, "num": 2})

    headers = {
        "X-API-KEY": os.getenv("SERPER_API_KEY"),
        "Content-Type": "application/json",
    }

    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(
                SERPER_URL, headers=headers, data=payload, timeout=30.0
            )
            response.raise_for_status()
            return response.json()
        except httpx.TimeoutException:
            return {"organic": []}

async def fetch_url(url: str):
    async with httpx.AsyncClient() as client:
        try:
            response = await client.get(url, timeout=30.0)
            soup = BeautifulSoup(response.text, "html.parser")
            text = soup.get_text()
            return text
        except httpx.TimeoutException:
            return "Timeout error"

@mcp.tool()
async def get_docs(query: str, library: str):
    """
    搜索给定查询和库的最新文档。
    支持 langchain、llama-index、autogen、agno、openai-agents-sdk、mcp-doc、camel-ai 和 crew-ai。

    参数:
    query: 要搜索的查询 (例如 "React Agent")
    library: 要搜索的库 (例如 "agno")

    返回:
    文档中的文本
    """
    if library not in docs_urls:
        raise ValueError(f"Library {library} not supported by this tool")

    query = f"site:{docs_urls[library]} {query}"
    results = await search_web(query)
    if len(results["organic"]) == 0:
        return "No results found"

    text = ""
    for result in results["organic"]:
        text += await fetch_url(result["link"])

    return text

@mcp.tool(description="查询符合输入姓名的人的电话号码")
def queryphone(a: str) -> dict:
    """
    查询某人的电话号码
    Args:
        a (str): 第一个字符串
    Returns:
       dict: 包含查询结果的字典
    
    """
    
    result: dict
    result = cxphone(a)
    return result 

def cxphone(name: str) -> dict:
    """
    查找并返回contacts.txt中指定名字的行
    Args:
        name (str): 要查找的名字
    Returns:
        dict: 包含名字和电话号码的字典列表
    """
    results = []
    with open('contacts.txt', 'r', encoding='utf-8') as file:
        for line in file:
            parts = line.strip().split('|')
            # 去掉所有的空格
            parts = [part.strip() for part in parts]
            if len(parts) >= 2 and name in parts[0]:  # 修改条件为包含name
                results.append({"name": parts[0], "phone": parts[1]})
    if results:
        return {"results": results}
    else:
        return {"error": f"未找到{name}的相关信息"}

#guy
documents = {}

# Helper Functions
def get_document_properties(doc_path: str) -> Dict[str, Any]:
    """Get properties of a Word document."""
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        core_props = doc.core_properties
        
        return {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "last_modified_by": core_props.last_modified_by or "",
            "revision": core_props.revision or 0,
            "page_count": len(doc.sections),
            "word_count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables)
        }
    except Exception as e:
        return {"error": f"Failed to get document properties: {str(e)}"}

def extract_document_text(doc_path: str) -> str:
    """Extract all text from a Word document."""
    if not os.path.exists(doc_path):
        return f"Document {doc_path} does not exist"
    
    try:
        doc = Document(doc_path)
        text = []
        
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text.append(paragraph.text)
        
        return "\n".join(text)
    except Exception as e:
        return f"Failed to extract text: {str(e)}"

def get_document_structure(doc_path: str) -> Dict[str, Any]:
    """Get the structure of a Word document."""
    if not os.path.exists(doc_path):
        return {"error": f"Document {doc_path} does not exist"}
    
    try:
        doc = Document(doc_path)
        structure = {
            "paragraphs": [],
            "tables": []
        }
        
        # Get paragraphs
        for i, para in enumerate(doc.paragraphs):
            structure["paragraphs"].append({
                "index": i,
                "text": para.text[:100] + ("..." if len(para.text) > 100 else ""),
                "style": para.style.name if para.style else "Normal"
            })
        
        # Get tables
        for i, table in enumerate(doc.tables):
            table_data = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "preview": []
            }
            
            # Get sample of table data
            max_rows = min(3, len(table.rows))
            for row_idx in range(max_rows):
                row_data = []
                max_cols = min(3, len(table.columns))
                for col_idx in range(max_cols):
                    try:
                        cell_text = table.cell(row_idx, col_idx).text
                        row_data.append(cell_text[:20] + ("..." if len(cell_text) > 20 else ""))
                    except IndexError:
                        row_data.append("N/A")
                table_data["preview"].append(row_data)
            
            structure["tables"].append(table_data)
        
        return structure
    except Exception as e:
        return {"error": f"Failed to get document structure: {str(e)}"}

def check_file_writeable(filepath: str) -> Tuple[bool, str]:
    """
    Check if a file can be written to.
    
    Args:
        filepath: Path to the file
        
    Returns:
        Tuple of (is_writeable, error_message)
    """
    # If file doesn't exist, check if directory is writeable
    if not os.path.exists(filepath):
        directory = os.path.dirname(filepath)
        if not os.path.exists(directory):
            return False, f"Directory {directory} does not exist"
        if not os.access(directory, os.W_OK):
            return False, f"Directory {directory} is not writeable"
        return True, ""
    
    # If file exists, check if it's writeable
    if not os.access(filepath, os.W_OK):
        return False, f"File {filepath} is not writeable (permission denied)"
    
    # Try to open the file for writing to see if it's locked
    try:
        with open(filepath, 'a'):
            pass
        return True, ""
    except IOError as e:
        return False, f"File {filepath} is not writeable: {str(e)}"
    except Exception as e:
        return False, f"Unknown error checking file permissions: {str(e)}"

def create_document_copy(source_path: str, dest_path: Optional[str] = None) -> Tuple[bool, str, Optional[str]]:
    """
    Create a copy of a document.
    
    Args:
        source_path: Path to the source document
        dest_path: Optional path for the new document. If not provided, will use source_path + '_copy.docx'
        
    Returns:
        Tuple of (success, message, new_filepath)
    """
    if not os.path.exists(source_path):
        return False, f"Source document {source_path} does not exist", None
    
    if not dest_path:
        # Generate a new filename if not provided
        base, ext = os.path.splitext(source_path)
        dest_path = f"{base}_copy{ext}"
    
    try:
        # Simple file copy
        shutil.copy2(source_path, dest_path)
        return True, f"Document copied to {dest_path}", dest_path
    except Exception as e:
        return False, f"Failed to copy document: {str(e)}", None

def ensure_heading_style(doc):
    """
    Ensure Heading styles exist in the document.
    
    Args:
        doc: Document object
    """
    for i in range(1, 10):  # Create Heading 1 through Heading 9
        style_name = f'Heading {i}'
        try:
            # Try to access the style to see if it exists
            style = doc.styles[style_name]
        except KeyError:
            # Create the style if it doesn't exist
            try:
                style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                if i == 1:
                    style.font.size = Pt(16)
                    style.font.bold = True
                elif i == 2:
                    style.font.size = Pt(14)
                    style.font.bold = True
                else:
                    style.font.size = Pt(12)
                    style.font.bold = True
            except Exception:
                # If style creation fails, we'll just use default formatting
                pass

def ensure_table_style(doc):
    """
    Ensure Table Grid style exists in the document.
    
    Args:
        doc: Document object
    """
    try:
        # Try to access the style to see if it exists
        style = doc.styles['Table Grid']
    except KeyError:
        # If style doesn't exist, we'll handle it at usage time
        pass

# MCP Tools
@mcp.tool()
async def create_document(filename: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """Create a new Word document with optional metadata.
    
    Args:
        filename: Name of the document to create (with or without .docx extension)
        title: Optional title for the document metadata
        author: Optional author for the document metadata
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot create document: {error_message}"
    
    try:
        doc = Document()
        
        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        # Ensure necessary styles exist
        ensure_heading_style(doc)
        ensure_table_style(doc)
        
        # Save the document
        doc.save(filename)
        
        return f"Document {filename} created successfully"
    except Exception as e:
        return f"Failed to create document: {str(e)}"

@mcp.tool()
async def add_heading(filename: str, text: str, level: int = 1) -> str:
    """Add a heading to a Word document.
    
    Args:
        filename: Path to the Word document
        text: Heading text
        level: Heading level (1-9, where 1 is the highest level)
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        
        # Ensure heading styles exist
        ensure_heading_style(doc)
        
        # Try to add heading with style
        try:
            heading = doc.add_heading(text, level=level)
            doc.save(filename)
            return f"Heading '{text}' (level {level}) added to {filename}"
        except Exception as style_error:
            # If style-based approach fails, use direct formatting
            paragraph = doc.add_paragraph(text)
            paragraph.style = doc.styles['Normal']
            run = paragraph.runs[0]
            run.bold = True
            # Adjust size based on heading level
            if level == 1:
                run.font.size = Pt(16)
            elif level == 2:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)
            
            doc.save(filename)
            return f"Heading '{text}' added to {filename} with direct formatting (style not available)"
    except Exception as e:
        return f"Failed to add heading: {str(e)}"

@mcp.tool()
async def add_paragraph(filename: str, text: str, style: Optional[str] = None) -> str:
    """Add a paragraph to a Word document.
    
    Args:
        filename: Path to the Word document
        text: Paragraph text
        style: Optional paragraph style name
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        paragraph = doc.add_paragraph(text)
        
        if style:
            try:
                paragraph.style = style
            except KeyError:
                # Style doesn't exist, use normal and report it
                paragraph.style = doc.styles['Normal']
                doc.save(filename)
                return f"Style '{style}' not found, paragraph added with default style to {filename}"
        
        doc.save(filename)
        return f"Paragraph added to {filename}"
    except Exception as e:
        return f"Failed to add paragraph: {str(e)}"

@mcp.tool()
async def add_table(filename: str, rows: int, cols: int, data: Optional[List[List[str]]] = None) -> str:
    """Add a table to a Word document.
    
    Args:
        filename: Path to the Word document
        rows: Number of rows in the table
        cols: Number of columns in the table
        data: Optional 2D array of data to fill the table
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        table = doc.add_table(rows=rows, cols=cols)
        
        # Try to set the table style
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If style doesn't exist, add basic borders
            # This is a simplified approach - complete border styling would require more code
            pass
        
        # Fill table with data if provided
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_text in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_text)
        
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        return f"Failed to add table: {str(e)}"

@mcp.tool()
async def add_picture(filename: str, image_path: str, width: Optional[float] = None) -> str:
    """Add an image to a Word document.
    
    Args:
        filename: Path to the Word document
        image_path: Path to the image file
        width: Optional width in inches (proportional scaling)
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    # Validate document existence
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Get absolute paths for better diagnostics
    abs_filename = os.path.abspath(filename)
    abs_image_path = os.path.abspath(image_path)
    
    # Validate image existence with improved error message
    if not os.path.exists(abs_image_path):
        return f"Image file not found: {abs_image_path}"
    
    # Check image file size
    try:
        image_size = os.path.getsize(abs_image_path) / 1024  # Size in KB
        if image_size <= 0:
            return f"Image file appears to be empty: {abs_image_path} (0 KB)"
    except Exception as size_error:
        return f"Error checking image file: {str(size_error)}"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(abs_filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(abs_filename)
        # Additional diagnostic info
        diagnostic = f"Attempting to add image ({abs_image_path}, {image_size:.2f} KB) to document ({abs_filename})"
        
        try:
            if width:
                doc.add_picture(abs_image_path, width=Inches(width))
            else:
                doc.add_picture(abs_image_path)
            doc.save(abs_filename)
            return f"Picture {image_path} added to {filename}"
        except Exception as inner_error:
            # More detailed error for the specific operation
            error_type = type(inner_error).__name__
            error_msg = str(inner_error)
            return f"Failed to add picture: {error_type} - {error_msg or 'No error details available'}\nDiagnostic info: {diagnostic}"
    except Exception as outer_error:
        # Fallback error handling
        error_type = type(outer_error).__name__
        error_msg = str(outer_error)
        return f"Document processing error: {error_type} - {error_msg or 'No error details available'}"

@mcp.tool()
async def get_document_info(filename: str) -> str:
    """Get information about a Word document.
    
    Args:
        filename: Path to the Word document
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        properties = get_document_properties(filename)
        return json.dumps(properties, indent=2)
    except Exception as e:
        return f"Failed to get document info: {str(e)}"

@mcp.tool()
async def get_document_text(filename: str) -> str:
    """Extract all text from a Word document.
    
    Args:
        filename: Path to the Word document
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    return extract_document_text(filename)

@mcp.tool()
async def get_document_outline(filename: str) -> str:
    """Get the structure of a Word document.
    
    Args:
        filename: Path to the Word document
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    structure = get_document_structure(filename)
    return json.dumps(structure, indent=2)

@mcp.tool()
async def list_available_documents(directory: str = ".") -> str:
    """List all .docx files in the specified directory.
    
    Args:
        directory: Directory to search for Word documents
    """
    try:
        if not os.path.exists(directory):
            return f"Directory {directory} does not exist"
        
        docx_files = [f for f in os.listdir(directory) if f.endswith('.docx')]
        
        if not docx_files:
            return f"No Word documents found in {directory}"
        
        result = f"Found {len(docx_files)} Word documents in {directory}:\n"
        for file in docx_files:
            file_path = os.path.join(directory, file)
            size = os.path.getsize(file_path) / 1024  # KB
            result += f"- {file} ({size:.2f} KB)\n"
        
        return result
    except Exception as e:
        return f"Failed to list documents: {str(e)}"

@mcp.tool()
async def copy_document(source_filename: str, destination_filename: Optional[str] = None) -> str:
    """Create a copy of a Word document.
    
    Args:
        source_filename: Path to the source document
        destination_filename: Optional path for the copy. If not provided, a default name will be generated.
    """
    if not source_filename.endswith('.docx'):
        source_filename += '.docx'
    
    if destination_filename and not destination_filename.endswith('.docx'):
        destination_filename += '.docx'
    
    success, message, new_path = create_document_copy(source_filename, destination_filename)
    if success:
        return message
    else:
        return f"Failed to copy document: {message}"

# Resources
@mcp.resource("docx:{path}")
async def document_resource(path: str) -> str:
    """Access Word document content."""
    if not path.endswith('.docx'):
        path += '.docx'
    
    if not os.path.exists(path):
        return f"Document {path} does not exist"
    
    return extract_document_text(path)
def find_paragraph_by_text(doc, text, partial_match=False):
    """
    Find paragraphs containing specific text.
    
    Args:
        doc: Document object
        text: Text to search for
        partial_match: If True, matches paragraphs containing the text; if False, matches exact text
        
    Returns:
        List of paragraph indices that match the criteria
    """
    matching_paragraphs = []
    
    for i, para in enumerate(doc.paragraphs):
        if partial_match and text in para.text:
            matching_paragraphs.append(i)
        elif not partial_match and para.text == text:
            matching_paragraphs.append(i)
            
    return matching_paragraphs

def find_and_replace_text(doc, old_text, new_text):
    """
    Find and replace text throughout the document.
    
    Args:
        doc: Document object
        old_text: Text to find
        new_text: Text to replace with
        
    Returns:
        Number of replacements made
    """
    count = 0
    
    # Search in paragraphs
    for para in doc.paragraphs:
        if old_text in para.text:
            for run in para.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
                    count += 1
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        for run in para.runs:
                            if old_text in run.text:
                                run.text = run.text.replace(old_text, new_text)
                                count += 1
    
    return count

def set_cell_border(cell, **kwargs):
    """
    Set cell border properties.
    
    Args:
        cell: The cell to modify
        **kwargs: Border properties (top, bottom, left, right, val, color)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Create border elements
    for key, value in kwargs.items():
        if key in ['top', 'left', 'bottom', 'right']:
            tag = 'w:{}'.format(key)
            
            element = OxmlElement(tag)
            element.set(qn('w:val'), kwargs.get('val', 'single'))
            element.set(qn('w:sz'), kwargs.get('sz', '4'))
            element.set(qn('w:space'), kwargs.get('space', '0'))
            element.set(qn('w:color'), kwargs.get('color', 'auto'))
            
            tcBorders = tcPr.first_child_found_in("w:tcBorders")
            if tcBorders is None:
                tcBorders = OxmlElement('w:tcBorders')
                tcPr.append(tcBorders)
                
            tcBorders.append(element)

def create_style(doc, style_name, style_type, base_style=None, font_properties=None, paragraph_properties=None):
    """
    Create a new style in the document.
    
    Args:
        doc: Document object
        style_name: Name for the new style
        style_type: Type of style (WD_STYLE_TYPE)
        base_style: Optional base style to inherit from
        font_properties: Dictionary of font properties (bold, italic, size, name, color)
        paragraph_properties: Dictionary of paragraph properties (alignment, spacing)
        
    Returns:
        The created style
    """
    try:
        # Check if style already exists
        style = doc.styles.get_by_id(style_name, WD_STYLE_TYPE.PARAGRAPH)
        return style
    except:
        # Create new style
        new_style = doc.styles.add_style(style_name, style_type)
        
        # Set base style if specified
        if base_style:
            new_style.base_style = doc.styles[base_style]
        
        # Set font properties
        if font_properties:
            font = new_style.font
            if 'bold' in font_properties:
                font.bold = font_properties['bold']
            if 'italic' in font_properties:
                font.italic = font_properties['italic']
            if 'size' in font_properties:
                font.size = Pt(font_properties['size'])
            if 'name' in font_properties:
                font.name = font_properties['name']
            if 'color' in font_properties:
                try:
                    # For RGB color
                    font.color.rgb = font_properties['color']
                except:
                    # For named color
                    font.color.theme_color = font_properties['color']
        
        # Set paragraph properties
        if paragraph_properties:
            if 'alignment' in paragraph_properties:
                new_style.paragraph_format.alignment = paragraph_properties['alignment']
            if 'spacing' in paragraph_properties:
                new_style.paragraph_format.line_spacing = paragraph_properties['spacing']
        
        return new_style

# Add these MCP tools to the existing set

@mcp.tool()
async def format_text(filename: str, paragraph_index: int, start_pos: int, end_pos: int, 
                     bold: Optional[bool] = None, italic: Optional[bool] = None, 
                     underline: Optional[bool] = None, color: Optional[str] = None,
                     font_size: Optional[int] = None, font_name: Optional[str] = None) -> str:
    """Format a specific range of text within a paragraph.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph (0-based)
        start_pos: Start position within the paragraph text
        end_pos: End position within the paragraph text
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        underline: Set text underlined (True/False)
        color: Text color (e.g., 'red', 'blue', etc.)
        font_size: Font size in points
        font_name: Font name/family
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        paragraph = doc.paragraphs[paragraph_index]
        text = paragraph.text
        
        # Validate text positions
        if start_pos < 0 or end_pos > len(text) or start_pos >= end_pos:
            return f"Invalid text positions. Paragraph has {len(text)} characters."
        
        # Get the text to format
        target_text = text[start_pos:end_pos]
        
        # Clear existing runs and create three runs: before, target, after
        for run in paragraph.runs:
            run.clear()
        
        # Add text before target
        if start_pos > 0:
            run_before = paragraph.add_run(text[:start_pos])
        
        # Add target text with formatting
        run_target = paragraph.add_run(target_text)
        if bold is not None:
            run_target.bold = bold
        if italic is not None:
            run_target.italic = italic
        if underline is not None:
            run_target.underline = underline
        if color:
            try:
                # Try to set color by name
                run_target.font.color.rgb = RGBColor.from_string(color)
            except:
                # If color name doesn't work, try predefined colors
                color_map = {
                    'red': WD_COLOR_INDEX.RED,
                    'blue': WD_COLOR_INDEX.BLUE,
                    'green': WD_COLOR_INDEX.GREEN,
                    'yellow': WD_COLOR_INDEX.YELLOW,
                    'black': WD_COLOR_INDEX.BLACK,
                }
                if color.lower() in color_map:
                    run_target.font.color.index = color_map[color.lower()]
        if font_size:
            run_target.font.size = Pt(font_size)
        if font_name:
            run_target.font.name = font_name
        
        # Add text after target
        if end_pos < len(text):
            run_after = paragraph.add_run(text[end_pos:])
        
        doc.save(filename)
        return f"Text '{target_text}' formatted successfully in paragraph {paragraph_index}."
    except Exception as e:
        return f"Failed to format text: {str(e)}"

@mcp.tool()
async def search_and_replace(filename: str, find_text: str, replace_text: str) -> str:
    """Search for text and replace all occurrences.
    
    Args:
        filename: Path to the Word document
        find_text: Text to search for
        replace_text: Text to replace with
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Perform find and replace
        count = find_and_replace_text(doc, find_text, replace_text)
        
        if count > 0:
            doc.save(filename)
            return f"Replaced {count} occurrence(s) of '{find_text}' with '{replace_text}'."
        else:
            return f"No occurrences of '{find_text}' found."
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"

# guy 得到个人总结
async def gy_get_grzj(question: str)-> str:
    client = OpenAI(
        base_url='https://api.siliconflow.cn/v1',
        api_key='sk-ausgzyjuyhyuaaizdxtzqltuimudowdrxwokgjrcgmebnwnm'
    )

    # 发送带有流式输出的请求
    response = client.chat.completions.create(
        # model="deepseek-ai/DeepSeek-V2.5",
        model="Qwen/Qwen2.5-72B-Instruct",
        messages=[
            {"role": "user", "content": question}
        ]
    )
    # 解析返回内容
    if response.choices:
        # 提取代码块（带Markdown格式检测）
        answer = response.choices[0].message.content
        return answer
    else:
        answer = "未获得有效响应"
        return answer

#生成公务员年度考核表
@mcp.tool()
async def gy_genetate_file(filename: str, find_text: str, replace_text: str) -> str:
    """生成指定姓名人员的公务员年度考核表.
    
    Args:
        filename: 指定的人员的姓名
        find_text: Text to search for
        replace_text: Text to replace with
    """
    await copy_document('tempbz.docx', filename +'.docx')
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # 使用正则表达式分割 find_text 和 replace_text
        find_texts = re.split(r'[,\s，]+', find_text.strip())
        replace_texts = re.split(r'[,\s，]+', replace_text.strip())
        
        # Ensure both lists have the same length
        if len(find_texts) != len(replace_texts):
            return f"Error: Number of find texts ({len(find_texts)}) does not match number of replace texts ({len(replace_texts)})."
        
        # Perform find and replace for each pair
        total_count = 0
        for find, replace in zip(find_texts, replace_texts):
            count = find_and_replace_text(doc, find.strip(), replace.strip())
            total_count += count
        
        # guy
        zj = await gy_get_grzj('写230字到234字的个人年度工作总结，身份是基层市场监管所一般工作人员，负责食品安全相关作')
        cleaned_zj = zj.replace("\r\n", "\n") 

        # 新增过滤逻辑
        # cleaned_zj = "\n".join([line for line in zj.split("\n") if line.strip() != ""])
        # cleaned_zj = cleaned_zj.replace('↓', '').replace('→', '').replace('←', '')  # 显式替换常见箭头
        cleaned_zj = re.sub(r'[\u2190-\u21FF]', '', cleaned_zj)  # 使用正则过滤Unicode箭头区字符
        # cleaned_zj = cleaned_zj.replace('\r', '')
        # 其他特殊符号处理
        cleaned_zj = cleaned_zj.translate(str.maketrans({
            '\u3000': ' ',   # 全角空格
            '\u00a0': ' ',   # 不间断空格
            '\u2028': '\n'   # 行分隔符转普通换行
        }))
        cleaned_zj = re.sub(r'[\r\n]+', ' ', cleaned_zj)
        print(cleaned_zj)
        find_and_replace_text(doc,'GRZJA'.strip(),cleaned_zj[:210].strip())
        find_and_replace_text(doc,'GRZJB'.strip(),cleaned_zj[210:].strip())
        
        if total_count > 0:
            doc.save(filename)
            return f"Replaced {total_count} occurrence(s) of specified texts."
        else:
            return f"No occurrences of specified texts found."
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"

@mcp.tool()
async def delete_paragraph(filename: str, paragraph_index: int) -> str:
    """Delete a paragraph from a document.
    
    Args:
        filename: Path to the Word document
        paragraph_index: Index of the paragraph to delete (0-based)
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        # Delete the paragraph (by removing its content and setting it empty)
        # Note: python-docx doesn't support true paragraph deletion, this is a workaround
        paragraph = doc.paragraphs[paragraph_index]
        p = paragraph._p
        p.getparent().remove(p)
        
        doc.save(filename)
        return f"Paragraph at index {paragraph_index} deleted successfully."
    except Exception as e:
        return f"Failed to delete paragraph: {str(e)}"

@mcp.tool()
async def create_custom_style(filename: str, style_name: str, 
                             bold: Optional[bool] = None, italic: Optional[bool] = None,
                             font_size: Optional[int] = None, font_name: Optional[str] = None,
                             color: Optional[str] = None, base_style: Optional[str] = None) -> str:
    """Create a custom style in the document.
    
    Args:
        filename: Path to the Word document
        style_name: Name for the new style
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        font_size: Font size in points
        font_name: Font name/family
        color: Text color (e.g., 'red', 'blue')
        base_style: Optional existing style to base this on
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Build font properties dictionary
        font_properties = {}
        if bold is not None:
            font_properties['bold'] = bold
        if italic is not None:
            font_properties['italic'] = italic
        if font_size is not None:
            font_properties['size'] = font_size
        if font_name is not None:
            font_properties['name'] = font_name
        if color is not None:
            font_properties['color'] = color
        
        # Create the style
        new_style = create_style(
            doc, 
            style_name, 
            WD_STYLE_TYPE.PARAGRAPH, 
            base_style=base_style,
            font_properties=font_properties
        )
        
        doc.save(filename)
        return f"Style '{style_name}' created successfully."
    except Exception as e:
        return f"Failed to create style: {str(e)}"

@mcp.tool()
async def format_table(filename: str, table_index: int, 
                      has_header_row: Optional[bool] = None,
                      border_style: Optional[str] = None,
                      shading: Optional[List[List[str]]] = None) -> str:
    """Format a table with borders, shading, and structure.
    
    Args:
        filename: Path to the Word document
        table_index: Index of the table (0-based)
        has_header_row: If True, formats the first row as a header
        border_style: Style for borders ('none', 'single', 'double', 'thick')
        shading: 2D list of cell background colors (by row and column)
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate table index
        if table_index < 0 or table_index >= len(doc.tables):
            return f"Invalid table index. Document has {len(doc.tables)} tables (0-{len(doc.tables)-1})."
        
        table = doc.tables[table_index]
        
        # Format header row if requested
        if has_header_row and table.rows:
            header_row = table.rows[0]
            for cell in header_row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.runs:
                        for run in paragraph.runs:
                            run.bold = True
        
        # Apply border style if specified
        if border_style:
            val_map = {
                'none': 'nil',
                'single': 'single',
                'double': 'double',
                'thick': 'thick'
            }
            val = val_map.get(border_style.lower(), 'single')
            
            # Apply to all cells
            for row in table.rows:
                for cell in row.cells:
                    set_cell_border(
                        cell,
                        top=True,
                        bottom=True,
                        left=True,
                        right=True,
                        val=val,
                        color="000000"
                    )
        
        # Apply cell shading if specified
        if shading:
            for i, row_colors in enumerate(shading):
                if i >= len(table.rows):
                    break
                for j, color in enumerate(row_colors):
                    if j >= len(table.rows[i].cells):
                        break
                    try:
                        # Apply shading to cell
                        cell = table.rows[i].cells[j]
                        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
                        cell._tc.get_or_add_tcPr().append(shading_elm)
                    except:
                        # Skip if color format is invalid
                        pass
        
        doc.save(filename)
        return f"Table at index {table_index} formatted successfully."
    except Exception as e:
        return f"Failed to format table: {str(e)}"

@mcp.tool()
async def add_page_break(filename: str) -> str:
    """Add a page break to the document.
    
    Args:
        filename: Path to the Word document
    """
    if not filename.endswith('.docx'):
        filename += '.docx'
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        doc.add_page_break()
        doc.save(filename)
        return f"Page break added to {filename}."
    except Exception as e:
        return f"Failed to add page break: {str(e)}"

## sse传输
def create_starlette_app(mcp_server: Server, *, debug: bool = False) -> Starlette:
    """Create a Starlette application that can serve the provided mcp server with SSE."""
    sse = SseServerTransport("/messages/")

    async def handle_sse(request: Request) -> None:
        async with sse.connect_sse(
                request.scope,
                request.receive,
                request._send,  # noqa: SLF001
        ) as (read_stream, write_stream):
            await mcp_server.run(
                read_stream,
                write_stream,
                mcp_server.create_initialization_options(),
            )

    return Starlette(
        debug=debug,
        routes=[
            Route("/sse", endpoint=handle_sse),
            Mount("/messages/", app=sse.handle_post_message),
        ],
    )

if __name__ == "__main__":
    mcp_server = mcp._mcp_server

    import argparse

    parser = argparse.ArgumentParser(description='Run MCP SSE-based server')
    parser.add_argument('--host', default='0.0.0.0', help='Host to bind to')
    parser.add_argument('--port', type=int, default=8020, help='Port to listen on')
    args = parser.parse_args()

    # Bind SSE request handling to MCP server
    starlette_app = create_starlette_app(mcp_server, debug=True)

    uvicorn.run(starlette_app, host=args.host, port=args.port)