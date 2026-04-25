"""
Document parsing utilities for memory/knowledge indexing.

This module normalizes multiple file formats into a unified text-plus-metadata
representation so downstream chunking can apply format-specific strategies.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional


SUPPORTED_DOCUMENT_EXTENSIONS = {
    ".md",
    ".markdown",
    ".txt",
    ".rst",
    ".pdf",
    ".docx",
    ".xlsx",
}


@dataclass
class ParsedDocument:
    """Normalized document representation for indexing."""

    content: str
    doc_type: str
    chunk_mode: str
    metadata: Dict[str, object] = field(default_factory=dict)


class DocumentParser:
    """Parse different office/document formats into normalized text."""

    def parse(self, file_path: Path, rel_path: str) -> ParsedDocument:
        suffix = file_path.suffix.lower()

        if suffix in {".md", ".markdown"}:
            return self._parse_markdown(file_path, rel_path)
        if suffix in {".txt", ".rst"}:
            return self._parse_text(file_path, rel_path)
        if suffix == ".pdf":
            return self._parse_pdf(file_path, rel_path)
        if suffix == ".docx":
            return self._parse_word(file_path, rel_path)
        if suffix == ".xlsx":
            return self._parse_excel(file_path, rel_path)

        raise ValueError(f"Unsupported document type for indexing: {suffix}")

    @staticmethod
    def is_supported(file_path: Path) -> bool:
        return file_path.suffix.lower() in SUPPORTED_DOCUMENT_EXTENSIONS

    def _parse_markdown(self, file_path: Path, rel_path: str) -> ParsedDocument:
        content = file_path.read_text(encoding="utf-8")
        title = self._extract_heading_title(content, fallback=file_path.stem)
        return ParsedDocument(
            content=content,
            doc_type="markdown",
            chunk_mode="markdown",
            metadata={
                "title": title,
                "parser": "markdown",
                "file_extension": file_path.suffix.lower(),
                "source_path": rel_path,
            },
        )

    def _parse_text(self, file_path: Path, rel_path: str) -> ParsedDocument:
        content = self._read_text_with_fallbacks(file_path)
        title = file_path.stem.replace("-", " ").replace("_", " ").strip() or file_path.name
        return ParsedDocument(
            content=content,
            doc_type="text",
            chunk_mode="plain_text",
            metadata={
                "title": title,
                "parser": "text",
                "file_extension": file_path.suffix.lower(),
                "source_path": rel_path,
            },
        )

    def _parse_pdf(self, file_path: Path, rel_path: str) -> ParsedDocument:
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf library not installed. Install with: pip install pypdf")

        reader = PdfReader(str(file_path))
        parts: List[str] = []
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = (page.extract_text() or "").strip()
            if page_text:
                parts.append(f"# Page {page_num}\n\n{page_text}")

        content = "\n\n".join(parts).strip()
        title = file_path.stem.replace("-", " ").replace("_", " ").strip() or file_path.name
        return ParsedDocument(
            content=content,
            doc_type="pdf",
            chunk_mode="pdf_pages",
            metadata={
                "title": title,
                "parser": "pdf",
                "page_count": len(reader.pages),
                "file_extension": file_path.suffix.lower(),
                "source_path": rel_path,
            },
        )

    def _parse_word(self, file_path: Path, rel_path: str) -> ParsedDocument:
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx library not installed. Install with: pip install python-docx")

        doc = Document(str(file_path))
        parts: List[str] = []
        title: Optional[str] = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = ""
            try:
                style_name = para.style.name or ""
            except Exception:
                style_name = ""

            normalized_style = style_name.lower()
            if normalized_style.startswith("title") and not title:
                title = text
                parts.append(f"# {text}")
            elif normalized_style.startswith("heading"):
                level = self._extract_heading_level(style_name)
                parts.append(f"{'#' * level} {text}")
            else:
                parts.append(text)

        for table_index, table in enumerate(doc.tables, start=1):
            parts.append(f"## Table {table_index}")
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))

        normalized = "\n\n".join(part for part in parts if part).strip()
        if not title:
            title = file_path.stem.replace("-", " ").replace("_", " ").strip() or file_path.name

        return ParsedDocument(
            content=normalized,
            doc_type="word",
            chunk_mode="word_sections",
            metadata={
                "title": title,
                "parser": "word",
                "file_extension": file_path.suffix.lower(),
                "source_path": rel_path,
            },
        )

    def _parse_excel(self, file_path: Path, rel_path: str) -> ParsedDocument:
        try:
            from openpyxl import load_workbook
        except ImportError:
            raise ImportError("openpyxl library not installed. Install with: pip install openpyxl")

        workbook = load_workbook(str(file_path), read_only=True, data_only=True)
        parts: List[str] = []
        non_empty_sheets = 0

        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            parts.append(f"# Sheet: {sheet_name}")

            rows = list(sheet.iter_rows(values_only=True))
            if not rows:
                continue

            header = None
            header_index = None
            for idx, row in enumerate(rows):
                cells = [str(cell).strip() if cell is not None else "" for cell in row]
                if any(cells):
                    header = cells
                    header_index = idx
                    break

            if header is None:
                continue

            non_empty_sheets += 1
            parts.append(" | ".join(header))

            for row_num, row in enumerate(rows[header_index + 1 :], start=header_index + 2):
                cells = [str(cell).strip() if cell is not None else "" for cell in row]
                if not any(cells):
                    continue
                row_pairs = []
                for col_index, cell_value in enumerate(cells):
                    column_name = header[col_index] if col_index < len(header) and header[col_index] else f"Column {col_index + 1}"
                    if cell_value:
                        row_pairs.append(f"{column_name}: {cell_value}")
                if row_pairs:
                    parts.append(f"- Row {row_num}: " + " | ".join(row_pairs))

        workbook.close()

        content = "\n\n".join(parts).strip()
        title = file_path.stem.replace("-", " ").replace("_", " ").strip() or file_path.name
        return ParsedDocument(
            content=content,
            doc_type="excel",
            chunk_mode="spreadsheet",
            metadata={
                "title": title,
                "parser": "excel",
                "sheet_count": len(workbook.sheetnames),
                "non_empty_sheets": non_empty_sheets,
                "file_extension": file_path.suffix.lower(),
                "source_path": rel_path,
            },
        )

    @staticmethod
    def _extract_heading_title(content: str, fallback: str) -> str:
        for line in content.splitlines():
            stripped = line.strip()
            if stripped.startswith("# "):
                return stripped[2:].strip()
        return fallback.replace("-", " ").replace("_", " ").strip() or fallback

    @staticmethod
    def _read_text_with_fallbacks(file_path: Path) -> str:
        encodings = ["utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"]
        for encoding in encodings:
            try:
                return file_path.read_text(encoding=encoding)
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise ValueError(f"Unable to decode text file: {file_path}")

    @staticmethod
    def _extract_heading_level(style_name: str) -> int:
        digits = "".join(ch for ch in style_name if ch.isdigit())
        if digits:
            try:
                level = int(digits)
                return min(max(level, 1), 6)
            except ValueError:
                return 2
        return 2
